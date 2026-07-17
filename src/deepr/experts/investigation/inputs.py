"""Zero-network input compilation for evidence-first investigations."""

from __future__ import annotations

import ipaddress
import mimetypes
import os
from collections.abc import Iterable, Sequence
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import SplitResult, urlsplit, urlunsplit

from deepr.experts.investigation.models import (
    INPUT_BUNDLE_KIND,
    INPUT_BUNDLE_SCHEMA_VERSION,
    InvestigationContractError,
    sha256_bytes,
    sha256_json,
    utc_now,
    validate_input_bundle,
)

DEFAULT_MAX_FILES = 256
DEFAULT_MAX_FILE_BYTES = 2 * 1024 * 1024
DEFAULT_MAX_TOTAL_BYTES = 32 * 1024 * 1024
DEFAULT_MAX_INLINE_BYTES = 256 * 1024
DEFAULT_MAX_EXTRACTED_BYTES = 512 * 1024
MAX_SCAN_ENTRIES = 4096

_TEXT_EXTENSIONS = frozenset(
    {
        ".cfg",
        ".conf",
        ".csv",
        ".css",
        ".graphql",
        ".htm",
        ".html",
        ".ini",
        ".java",
        ".js",
        ".json",
        ".jsonl",
        ".jsx",
        ".md",
        ".mdx",
        ".py",
        ".rst",
        ".sh",
        ".sql",
        ".toml",
        ".ts",
        ".tsx",
        ".txt",
        ".xml",
        ".yaml",
        ".yml",
    }
)
_TEXT_FILENAMES = frozenset({"dockerfile", "license", "makefile", "readme"})
_DOCUMENT_EXTENSIONS = frozenset({".docx"})
_EXCLUDED_DIR_NAMES = frozenset(
    {
        ".aws",
        ".azure",
        ".git",
        ".hg",
        ".idea",
        ".mypy_cache",
        ".pytest_cache",
        ".ruff_cache",
        ".ssh",
        ".svn",
        ".venv",
        ".vscode",
        "__pycache__",
        "build",
        "coverage",
        "dist",
        "node_modules",
        "target",
        "venv",
    }
)
_SECRET_EXACT_NAMES = frozenset(
    {
        ".env",
        ".env.local",
        ".netrc",
        "credentials.json",
        "id_dsa",
        "id_ed25519",
        "id_rsa",
        "known_hosts",
        "secrets.json",
    }
)
_SECRET_SUFFIXES = (".key", ".p12", ".pfx", ".pem")


@dataclass(frozen=True)
class InputLimits:
    """Hard local-input bounds included in the immutable plan."""

    max_files: int = DEFAULT_MAX_FILES
    max_file_bytes: int = DEFAULT_MAX_FILE_BYTES
    max_total_bytes: int = DEFAULT_MAX_TOTAL_BYTES
    max_inline_bytes: int = DEFAULT_MAX_INLINE_BYTES
    max_extracted_bytes: int = DEFAULT_MAX_EXTRACTED_BYTES

    def validated(self) -> InputLimits:
        values = {
            "max_files": self.max_files,
            "max_file_bytes": self.max_file_bytes,
            "max_total_bytes": self.max_total_bytes,
            "max_inline_bytes": self.max_inline_bytes,
            "max_extracted_bytes": self.max_extracted_bytes,
        }
        for name, value in values.items():
            if isinstance(value, bool) or not isinstance(value, int) or value < 1:
                raise InvestigationContractError(f"{name} must be a positive integer")
        if self.max_file_bytes > self.max_total_bytes:
            raise InvestigationContractError("max_file_bytes cannot exceed max_total_bytes")
        return self

    def to_dict(self) -> dict[str, int]:
        return {
            "max_files": self.max_files,
            "max_file_bytes": self.max_file_bytes,
            "max_total_bytes": self.max_total_bytes,
            "max_inline_bytes": self.max_inline_bytes,
            "max_extracted_bytes": self.max_extracted_bytes,
        }


def _inside_root(path: Path, root: Path) -> bool:
    try:
        path.relative_to(root)
    except ValueError:
        return False
    return True


def _relative_display(path: Path, root: Path) -> str:
    return path.relative_to(root).as_posix()


def _hidden(relative: Path) -> bool:
    return any(part.startswith(".") for part in relative.parts)


def _secret_bearing(relative: Path) -> bool:
    name = relative.name.casefold()
    if name in _SECRET_EXACT_NAMES or name.startswith(".env."):
        return True
    return name.endswith(_SECRET_SUFFIXES)


def _supported_file(path: Path) -> tuple[bool, str, str]:
    suffix = path.suffix.casefold()
    filename = path.name.casefold()
    media_type = mimetypes.guess_type(path.name)[0] or "text/plain"
    if suffix in _TEXT_EXTENSIONS or filename in _TEXT_FILENAMES:
        return True, media_type, "utf8-text-v1"
    if suffix in _DOCUMENT_EXTENSIONS:
        return True, media_type, "docx-text-v1"
    return False, media_type, "unsupported"


def _exclusion(path: str, reason: str, *, origin: str, detail: str) -> dict[str, str]:
    return {"path": path, "reason": reason, "origin": origin, "detail": detail}


def _normalize_url(raw_url: str) -> str:
    value = raw_url.strip()
    if not value:
        raise InvestigationContractError("URL inputs must be non-empty")
    try:
        parts = urlsplit(value)
        port = parts.port
    except ValueError as exc:
        raise InvestigationContractError(f"invalid URL: {exc}") from exc
    if parts.scheme.casefold() not in {"http", "https"} or not parts.hostname:
        raise InvestigationContractError("URL inputs require an http or https host")
    if parts.username is not None or parts.password is not None:
        raise InvestigationContractError("URL inputs cannot contain credentials")
    host = parts.hostname.casefold().rstrip(".")
    if host == "localhost" or host.endswith(".localhost"):
        raise InvestigationContractError("URL inputs cannot target localhost")
    try:
        literal = ipaddress.ip_address(host.strip("[]"))
    except ValueError:
        literal = None
    if literal is not None and not literal.is_global:
        raise InvestigationContractError("URL inputs cannot target non-public IP literals")
    netloc = f"[{host}]" if literal is not None and literal.version == 6 else host
    if port is not None:
        netloc = f"{netloc}:{port}"
    normalized = SplitResult(parts.scheme.casefold(), netloc, parts.path or "/", parts.query, "")
    return urlunsplit(normalized)


def _resolve_candidate(raw_path: str | Path, *, root: Path) -> tuple[Path, str]:
    candidate = Path(raw_path)
    if not candidate.is_absolute():
        candidate = root / candidate
    lexical = Path(os.path.abspath(candidate))
    if not _inside_root(lexical, root):
        raise InvestigationContractError(f"input path escapes the input root: {candidate}")
    current = root
    for part in lexical.relative_to(root).parts:
        current /= part
        if current.is_symlink():
            raise InvestigationContractError(f"input path traverses a symbolic link: {candidate}")
    try:
        resolved = candidate.resolve(strict=True)
    except (OSError, RuntimeError) as exc:
        raise InvestigationContractError(f"input path is unavailable: {candidate}") from exc
    if not _inside_root(resolved, root):
        raise InvestigationContractError(f"input path escapes the input root: {candidate}")
    return resolved, _relative_display(resolved, root)


def _folder_candidates(folder: Path, *, root: Path, origin: str) -> tuple[list[Path], list[dict[str, str]]]:
    candidates: list[Path] = []
    exclusions: list[dict[str, str]] = []
    scanned = 0
    for current_text, dir_names, file_names in os.walk(folder, topdown=True, followlinks=False):
        current = Path(current_text)
        kept_dirs: list[str] = []
        for name in sorted(dir_names, key=str.casefold):
            scanned += 1
            child = current / name
            display = _relative_display(child, root)
            relative = Path(display)
            if scanned > MAX_SCAN_ENTRIES:
                raise InvestigationContractError(f"folder scan exceeds {MAX_SCAN_ENTRIES} entries")
            if child.is_symlink():
                exclusions.append(
                    _exclusion(display, "symlink", origin=origin, detail="symbolic links are not followed")
                )
            elif _hidden(relative) or name.casefold() in _EXCLUDED_DIR_NAMES:
                exclusions.append(
                    _exclusion(
                        display, "excluded_subtree", origin=origin, detail="hidden, metadata, cache, or build directory"
                    )
                )
            else:
                kept_dirs.append(name)
        dir_names[:] = kept_dirs
        for name in sorted(file_names, key=str.casefold):
            scanned += 1
            if scanned > MAX_SCAN_ENTRIES:
                raise InvestigationContractError(f"folder scan exceeds {MAX_SCAN_ENTRIES} entries")
            candidates.append(current / name)
    return candidates, exclusions


def _file_item(path: Path, *, root: Path, origin: str) -> tuple[dict[str, Any] | None, dict[str, str] | None]:
    display = _relative_display(path, root)
    relative = Path(display)
    if path.is_symlink():
        return None, _exclusion(display, "symlink", origin=origin, detail="symbolic links are not read")
    if _hidden(relative):
        return None, _exclusion(display, "hidden_path", origin=origin, detail="hidden paths are excluded")
    if _secret_bearing(relative):
        return None, _exclusion(
            display, "secret_bearing_path", origin=origin, detail="credential and key paths are excluded"
        )
    supported, media_type, extractor = _supported_file(path)
    if not supported:
        return None, _exclusion(display, "unsupported_type", origin=origin, detail="no admitted text extractor")
    try:
        raw = path.read_bytes()
    except OSError as exc:
        return None, _exclusion(display, "unreadable", origin=origin, detail=type(exc).__name__)
    if extractor == "utf8-text-v1" and b"\x00" in raw[:8192]:
        return None, _exclusion(display, "binary_content", origin=origin, detail="NUL bytes found in text input")
    return (
        {
            "input_type": "file",
            "source_class": "caller_supplied_file",
            "display_path": display,
            "origin": origin,
            "media_type": media_type,
            "byte_size": len(raw),
            "content_sha256": sha256_bytes(raw),
            "extraction_status": "ready",
            "extractor": extractor,
        },
        None,
    )


def _resolve_input_root(input_root: str | Path) -> Path:
    try:
        root = Path(input_root).resolve(strict=True)
    except (OSError, RuntimeError) as exc:
        raise InvestigationContractError(f"input root is unavailable: {input_root}") from exc
    if not root.is_dir():
        raise InvestigationContractError("input root must be a directory")
    return root


def _compile_inline_items(texts: Sequence[str], limits: InputLimits) -> tuple[list[dict[str, Any]], int]:
    items: list[dict[str, Any]] = []
    total_bytes = 0
    for index, text in enumerate(texts, start=1):
        if not isinstance(text, str) or not text.strip():
            raise InvestigationContractError("inline text inputs must be non-empty strings")
        raw = text.encode("utf-8")
        if len(raw) > limits.max_inline_bytes:
            raise InvestigationContractError(f"inline text {index} exceeds max_inline_bytes")
        total_bytes += len(raw)
        items.append(
            {
                "input_type": "inline_text",
                "source_class": "caller_supplied",
                "label": f"inline-{index}",
                "media_type": "text/plain; charset=utf-8",
                "byte_size": len(raw),
                "content_sha256": sha256_bytes(raw),
                "extraction_status": "ready",
                "content": text,
            }
        )
    return items, total_bytes


def _compile_url_items(urls: Sequence[str]) -> tuple[list[dict[str, Any]], set[str]]:
    items: list[dict[str, Any]] = []
    normalized_urls: set[str] = set()
    for raw_url in urls:
        normalized = _normalize_url(raw_url)
        if normalized in normalized_urls:
            continue
        normalized_urls.add(normalized)
        items.append(
            {
                "input_type": "url",
                "source_class": "requested_external_source",
                "url": normalized,
                "media_type": "application/octet-stream",
                "byte_size": 0,
                "content_sha256": sha256_bytes(normalized.encode("utf-8")),
                "extraction_status": "pending_fetch",
            }
        )
    return items, normalized_urls


def _collect_file_candidates(
    files: Sequence[str | Path],
    folders: Sequence[str | Path],
    *,
    root: Path,
) -> tuple[list[tuple[Path, str]], list[dict[str, str]]]:
    candidates: list[tuple[Path, str]] = []
    exclusions: list[dict[str, str]] = []
    for raw_file in files:
        resolved, display = _resolve_candidate(raw_file, root=root)
        if not resolved.is_file():
            raise InvestigationContractError(f"file input is not a regular file: {display}")
        candidates.append((resolved, f"file:{display}"))
    for raw_folder in folders:
        resolved, display = _resolve_candidate(raw_folder, root=root)
        if not resolved.is_dir():
            raise InvestigationContractError(f"folder input is not a directory: {display}")
        folder_candidates, folder_exclusions = _folder_candidates(resolved, root=root, origin=f"folder:{display}")
        exclusions.extend(folder_exclusions)
        candidates.extend((candidate, f"folder:{display}") for candidate in folder_candidates)
    return candidates, exclusions


def _file_limit_exclusion(
    item: dict[str, Any],
    *,
    origin: str,
    limits: InputLimits,
    included_files: int,
    total_bytes: int,
) -> dict[str, str] | None:
    size = int(item["byte_size"])
    if size > limits.max_file_bytes:
        return _exclusion(item["display_path"], "oversized_file", origin=origin, detail=f"{size} bytes exceeds limit")
    if included_files >= limits.max_files:
        return _exclusion(
            item["display_path"], "file_count_limit", origin=origin, detail="maximum included files reached"
        )
    if total_bytes + size > limits.max_total_bytes:
        return _exclusion(
            item["display_path"],
            "aggregate_size_limit",
            origin=origin,
            detail="maximum aggregate bytes reached",
        )
    return None


def _compile_file_items(
    candidates: list[tuple[Path, str]],
    *,
    root: Path,
    limits: InputLimits,
    initial_bytes: int,
) -> tuple[list[dict[str, Any]], list[dict[str, str]], int, int]:
    items: list[dict[str, Any]] = []
    exclusions: list[dict[str, str]] = []
    seen_paths: set[Path] = set()
    total_bytes = initial_bytes
    included_files = 0
    ordered = sorted(candidates, key=lambda value: _relative_display(value[0], root).casefold())
    for candidate, origin in ordered:
        if candidate.is_symlink():
            display = _relative_display(candidate, root)
            exclusions.append(_exclusion(display, "symlink", origin=origin, detail="symbolic links are not read"))
            continue
        try:
            resolved = candidate.resolve(strict=True)
        except (OSError, RuntimeError):
            exclusions.append(
                _exclusion(str(candidate), "unavailable", origin=origin, detail="path changed during preview")
            )
            continue
        if not _inside_root(resolved, root) or resolved in seen_paths:
            continue
        seen_paths.add(resolved)
        item, exclusion = _file_item(resolved, root=root, origin=origin)
        if exclusion is not None:
            exclusions.append(exclusion)
            continue
        if item is None:
            raise InvestigationContractError("file input produced neither an item nor an exclusion")
        limit_exclusion = _file_limit_exclusion(
            item,
            origin=origin,
            limits=limits,
            included_files=included_files,
            total_bytes=total_bytes,
        )
        if limit_exclusion is not None:
            exclusions.append(limit_exclusion)
            continue
        items.append(item)
        total_bytes += int(item["byte_size"])
        included_files += 1
    return items, exclusions, total_bytes, included_files


def compile_input_bundle(
    *,
    input_root: str | Path,
    inline_texts: Sequence[str] = (),
    urls: Sequence[str] = (),
    files: Sequence[str | Path] = (),
    folders: Sequence[str | Path] = (),
    limits: InputLimits | None = None,
    created_at: str | None = None,
) -> dict[str, Any]:
    """Hash caller inputs without fetching URLs, invoking models, or writing files."""
    effective_limits = (limits or InputLimits()).validated()
    root = _resolve_input_root(input_root)
    inline_items, total_bytes = _compile_inline_items(inline_texts, effective_limits)
    url_items, normalized_urls = _compile_url_items(urls)
    file_candidates, folder_exclusions = _collect_file_candidates(files, folders, root=root)
    file_items, file_exclusions, total_bytes, included_files = _compile_file_items(
        file_candidates,
        root=root,
        limits=effective_limits,
        initial_bytes=total_bytes,
    )
    items = [*inline_items, *url_items, *file_items]
    exclusions = [*folder_exclusions, *file_exclusions]

    if total_bytes > effective_limits.max_total_bytes:
        raise InvestigationContractError("inline inputs exceed max_total_bytes")
    for index, item in enumerate(items, start=1):
        item["input_id"] = f"input-{index:04d}"
    exclusions.sort(key=lambda item: (item["path"].casefold(), item["reason"], item["origin"]))
    material: dict[str, Any] = {
        "schema_version": INPUT_BUNDLE_SCHEMA_VERSION,
        "kind": INPUT_BUNDLE_KIND,
        "root": str(root),
        "created_at": created_at or utc_now(),
        "limits": effective_limits.to_dict(),
        "items": items,
        "exclusions": exclusions,
        "summary": {
            "included_items": len(items),
            "included_files": included_files,
            "requested_urls": len(normalized_urls),
            "excluded_paths": len(exclusions),
            "local_input_bytes": total_bytes,
        },
    }
    material["bundle_sha256"] = sha256_json(material)
    return validate_input_bundle(material)


def _decode_text(raw: bytes, *, display_path: str) -> str:
    try:
        return raw.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise InvestigationContractError(f"text input is not UTF-8: {display_path}") from exc


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise InvestigationContractError("DOCX extraction requires the docs extra") from exc
    try:
        document = Document(str(path))
    except (OSError, ValueError) as exc:
        raise InvestigationContractError(f"DOCX extraction failed: {path.name}") from exc
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            value = " | ".join(cell.text.strip() for cell in row.cells)
            if value.strip(" |"):
                paragraphs.append(value)
    return "\n".join(paragraphs)


def materialize_input_context(bundle: dict[str, Any]) -> list[dict[str, str]]:
    """Verify frozen inputs and return bounded untrusted text for model packets."""
    validated = validate_input_bundle(bundle)
    root = Path(validated["root"])
    max_extracted = int(validated["limits"]["max_extracted_bytes"])
    remaining = max_extracted
    context: list[dict[str, str]] = []
    for item in validated["items"]:
        input_type = item["input_type"]
        if input_type == "url":
            continue
        if input_type == "inline_text":
            text = str(item["content"])
            label = str(item["label"])
            source_class = "caller_supplied"
        else:
            display_path = str(item["display_path"])
            path, confirmed_display = _resolve_candidate(display_path, root=root)
            if confirmed_display != display_path or not path.is_file():
                raise InvestigationContractError(f"frozen input path changed: {display_path}")
            raw = path.read_bytes()
            if sha256_bytes(raw) != item["content_sha256"]:
                raise InvestigationContractError(f"frozen input content changed: {display_path}")
            extractor = item["extractor"]
            text = _decode_text(raw, display_path=display_path) if extractor == "utf8-text-v1" else _extract_docx(path)
            label = display_path
            source_class = "caller_supplied_file"
        encoded = text.encode("utf-8")
        if remaining <= 0:
            break
        excerpt = encoded[:remaining].decode("utf-8", errors="ignore")
        remaining -= len(excerpt.encode("utf-8"))
        context.append(
            {
                "ref": str(item["input_id"]),
                "label": label,
                "source_class": source_class,
                "text": excerpt,
            }
        )
    return context


def requested_urls(bundle: dict[str, Any]) -> tuple[str, ...]:
    """Return the normalized URL requests without fetching them."""
    validated = validate_input_bundle(bundle)
    return tuple(str(item["url"]) for item in validated["items"] if item["input_type"] == "url")


def included_paths(bundle: dict[str, Any]) -> Iterable[str]:
    """Yield frozen relative file paths for display and inspection."""
    validated = validate_input_bundle(bundle)
    return (str(item["display_path"]) for item in validated["items"] if item["input_type"] == "file")


__all__ = [
    "DEFAULT_MAX_EXTRACTED_BYTES",
    "DEFAULT_MAX_FILES",
    "DEFAULT_MAX_FILE_BYTES",
    "DEFAULT_MAX_INLINE_BYTES",
    "DEFAULT_MAX_TOTAL_BYTES",
    "InputLimits",
    "compile_input_bundle",
    "included_paths",
    "materialize_input_context",
    "requested_urls",
]
