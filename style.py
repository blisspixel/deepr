import re
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Added a new style constant for Heading 3 ---
STYLE_NORMAL = "Normal"
STYLE_HEADING_1 = "Heading 1"
STYLE_HEADING_2 = "Heading 2"
STYLE_HEADING_3 = "Heading 3" # New style for numbered sections
STYLE_BULLET = "List Bullet"
STYLE_NUMBERED = "List Numbered"
STYLE_BLOCKQUOTE = "Blockquote"
STYLE_CODE = "Code"

# --- All helper functions remain the same ---
def safe_get_style(doc, name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    r.font.underline = True
    return hyperlink

def add_formatted_text(paragraph, text_line):
    pattern = re.compile(r'(\[.*?\]\(.*?\))|(\*\*.*?\*\*)|(\*.*?\*)|(https?://[^\s)]+)')
    cursor = 0
    for match in pattern.finditer(text_line):
        start, end = match.span()
        if start > cursor:
            paragraph.add_run(text_line[cursor:start])
        full_match = match.group(0)
        if full_match.startswith('['):
            link_text = full_match[1:full_match.find(']')]
            link_url = full_match[full_match.find('(') + 1:-1]
            add_hyperlink(paragraph, link_url, link_text)
        elif full_match.startswith('**'):
            run = paragraph.add_run(full_match[2:-2])
            run.bold = True
        elif full_match.startswith('*'):
            run = paragraph.add_run(full_match[1:-1])
            run.italic = True
        elif full_match.startswith('http'):
            add_hyperlink(paragraph, full_match, full_match)
        cursor = end
    if cursor < len(text_line):
        paragraph.add_run(text_line[cursor:])

# --- UPDATED TO INCLUDE HEADING 3 DEFINITION ---
def apply_basic_styles(doc):
    style = safe_get_style(doc, STYLE_NORMAL); font = style.font; font.name = "Aptos Body"; font.size = Pt(12)
    heading_1_style = safe_get_style(doc, STYLE_HEADING_1); font = heading_1_style.font; font.name = "Aptos Display"; font.size = Pt(16); font.bold = True
    heading_2_style = safe_get_style(doc, STYLE_HEADING_2); font = heading_2_style.font; font.name = "Aptos Display"; font.size = Pt(14); font.bold = True
    
    # Define the new Heading 3 style
    heading_3_style = safe_get_style(doc, STYLE_HEADING_3); font = heading_3_style.font; font.name = "Aptos Display"; font.size = Pt(13); font.bold = True

    bullet_style = safe_get_style(doc, STYLE_BULLET); font = bullet_style.font; font.name = "Aptos Body"; font.size = Pt(12)
    blockquote_style = safe_get_style(doc, STYLE_BLOCKQUOTE); font = blockquote_style.font; font.name = "Aptos Body"; font.size = Pt(12); font.italic = True

def format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_after=6, spacing_before=0, indent=0):
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.space_after = Pt(spacing_after)
    pf.space_before = Pt(spacing_before)
    pf.line_spacing = 1.08
    pf.left_indent = Inches(indent)
    pf.right_indent = Inches(0)

# --- UPDATED MAIN DOCUMENT PROCESSING FUNCTION ---
def apply_styles_to_doc(doc, full_text):
    apply_basic_styles(doc)
    lines = full_text.split("\n")

    for line in lines:
        stripped_line = line.strip()

        # --- ADDED LOGIC TO SKIP SEPARATOR LINES AND EMPTY LINES ---
        if not stripped_line or stripped_line == "---":
            continue

        # Handle Markdown-style headings
        if stripped_line.startswith("# "):
            para = doc.add_paragraph(stripped_line[2:], style=STYLE_HEADING_1)
            format_paragraph(para, spacing_before=12, spacing_after=4)
            continue
        if stripped_line.startswith("## "):
            para = doc.add_paragraph(stripped_line[3:], style=STYLE_HEADING_2)
            format_paragraph(para, spacing_before=10, spacing_after=3)
            continue

        # --- NEW: Handle numbered lines as Heading 3 ---
        # This rule now takes precedence over numbered lists.
        match = re.match(r'^\d+\.\s+(.*)', stripped_line)
        if match:
            # Strip the number and treat the rest as a Heading 3
            heading_text = match.group(1).strip()
            para = doc.add_paragraph(style=STYLE_HEADING_3)
            # Use the formatter to handle any bold/italic text within the heading
            add_formatted_text(para, heading_text)
            format_paragraph(para, spacing_before=10, spacing_after=3)
            continue
            
        # Case 1: A line that acts as a BOLD HEADING for a list (e.g., "- **Title:**")
        if stripped_line.startswith("- **") and stripped_line.endswith(":**"):
            para = doc.add_paragraph()
            format_paragraph(para, spacing_before=8, spacing_after=2, indent=0.25) # Indent to align with list
            text_to_parse = stripped_line[2:]
            add_formatted_text(para, text_to_parse)

        # Case 2: A regular bullet point
        elif stripped_line.startswith(("- ", "* ")):
            para = doc.add_paragraph(style=STYLE_BULLET)
            format_paragraph(para, indent=0.5) # Deeper indent for sub-items
            text_to_parse = stripped_line[2:]
            add_formatted_text(para, text_to_parse)

        # Case 3: A blockquote
        elif stripped_line.startswith("> "):
            para = doc.add_paragraph(style=STYLE_BLOCKQUOTE)
            format_paragraph(para, indent=0.5)
            text_to_parse = stripped_line[2:]
            add_formatted_text(para, text_to_parse)

        # Case 4: All other lines are treated as normal paragraphs
        else:
            para = doc.add_paragraph(style=STYLE_NORMAL)
            format_paragraph(para)
            text_to_parse = stripped_line
            add_formatted_text(para, text_to_parse)
            
    return doc